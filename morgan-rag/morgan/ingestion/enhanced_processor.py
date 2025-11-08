"""Production-quality document processor with streaming and async processing."""

import asyncio
import hashlib
import logging
import mimetypes
import time
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, AsyncIterator, Dict, List, Optional, Set

import aiofiles
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class ChunkingStrategy(Enum):
    """Document chunking strategies."""
    FIXED = "fixed"
    SEMANTIC = "semantic"
    RECURSIVE = "recursive"


@dataclass
class ProcessingConfig:
    """Document processing configuration."""
    chunk_size: int = 1000
    chunk_overlap: int = 200
    chunking_strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE
    max_file_size_mb: int = 100
    supported_extensions: Set[str] = field(default_factory=lambda: {
        '.txt', '.md', '.pdf', '.docx', '.html', '.htm',
        '.py', '.js', '.ts', '.java', '.cpp', '.c', '.h',
        '.json', '.yaml', '.yml', '.xml', '.csv'
    })
    max_concurrent_files: int = 10
    stream_buffer_size: int = 50
    enable_metadata_extraction: bool = True


@dataclass
class DocumentChunk:
    """Processed document chunk."""
    chunk_id: str
    document_id: str
    content: str
    chunk_index: int
    total_chunks: int
    metadata: Dict[str, Any]
    char_start: int
    char_end: int


@dataclass
class ProcessedDocument:
    """Processed document with chunks."""
    document_id: str
    file_path: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processing_time: float
    error: Optional[str] = None


class TextExtractor:
    """Async text extraction from various file formats."""

    @staticmethod
    async def extract_text(file_path: Path) -> str:
        """Extract text from file based on extension.

        Args:
            file_path: Path to file.

        Returns:
            Extracted text content.

        Raises:
            ValueError: If file type not supported.
        """
        ext = file_path.suffix.lower()

        if ext in {'.txt', '.md', '.py', '.js', '.ts', '.java', '.cpp', '.c', '.h', '.json', '.yaml', '.yml', '.xml'}:
            return await TextExtractor._extract_plain_text(file_path)
        elif ext == '.pdf':
            return await TextExtractor._extract_pdf(file_path)
        elif ext == '.docx':
            return await TextExtractor._extract_docx(file_path)
        elif ext in {'.html', '.htm'}:
            return await TextExtractor._extract_html(file_path)
        elif ext == '.csv':
            return await TextExtractor._extract_csv(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    async def _extract_plain_text(file_path: Path) -> str:
        """Extract plain text file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return await f.read()

    @staticmethod
    async def _extract_pdf(file_path: Path) -> str:
        """Extract text from PDF."""
        # Run in thread pool to avoid blocking
        loop = asyncio.get_event_loop()

        def _sync_extract():
            try:
                import pypdf
                text_parts = []
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    for page in reader.pages:
                        text_parts.append(page.extract_text())
                return '\n\n'.join(text_parts)
            except Exception as e:
                logger.error(f"PDF extraction failed: {e}")
                return ""

        return await loop.run_in_executor(None, _sync_extract)

    @staticmethod
    async def _extract_docx(file_path: Path) -> str:
        """Extract text from DOCX."""
        loop = asyncio.get_event_loop()

        def _sync_extract():
            try:
                from docx import Document
                doc = Document(file_path)
                return '\n\n'.join(para.text for para in doc.paragraphs)
            except Exception as e:
                logger.error(f"DOCX extraction failed: {e}")
                return ""

        return await loop.run_in_executor(None, _sync_extract)

    @staticmethod
    async def _extract_html(file_path: Path) -> str:
        """Extract text from HTML."""
        loop = asyncio.get_event_loop()

        def _sync_extract():
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                    # Remove script and style elements
                    for script in soup(["script", "style"]):
                        script.decompose()
                    return soup.get_text(separator='\n', strip=True)
            except Exception as e:
                logger.error(f"HTML extraction failed: {e}")
                return ""

        return await loop.run_in_executor(None, _sync_extract)

    @staticmethod
    async def _extract_csv(file_path: Path) -> str:
        """Extract text from CSV."""
        loop = asyncio.get_event_loop()

        def _sync_extract():
            try:
                import csv
                rows = []
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        rows.append(' | '.join(row))
                return '\n'.join(rows)
            except Exception as e:
                logger.error(f"CSV extraction failed: {e}")
                return ""

        return await loop.run_in_executor(None, _sync_extract)


class ChunkGenerator:
    """Generate document chunks using various strategies."""

    def __init__(self, config: ProcessingConfig):
        self.config = config
        self._splitters: Dict[ChunkingStrategy, Any] = {}

    def _get_splitter(self, strategy: ChunkingStrategy):
        """Get or create text splitter for strategy."""
        if strategy not in self._splitters:
            if strategy == ChunkingStrategy.RECURSIVE:
                self._splitters[strategy] = RecursiveCharacterTextSplitter(
                    chunk_size=self.config.chunk_size,
                    chunk_overlap=self.config.chunk_overlap,
                    separators=["\n\n", "\n", ". ", " ", ""],
                    length_function=len,
                )
            elif strategy == ChunkingStrategy.FIXED:
                self._splitters[strategy] = RecursiveCharacterTextSplitter(
                    chunk_size=self.config.chunk_size,
                    chunk_overlap=self.config.chunk_overlap,
                    separators=[""],  # Character-based only
                    length_function=len,
                )
            else:
                # Default to recursive
                self._splitters[strategy] = RecursiveCharacterTextSplitter(
                    chunk_size=self.config.chunk_size,
                    chunk_overlap=self.config.chunk_overlap,
                )

        return self._splitters[strategy]

    async def generate_chunks(
        self,
        text: str,
        document_id: str,
        metadata: Dict[str, Any],
    ) -> List[DocumentChunk]:
        """Generate chunks from text.

        Args:
            text: Input text.
            document_id: Document identifier.
            metadata: Document metadata.

        Returns:
            List of document chunks.
        """
        if not text.strip():
            return []

        # Run splitting in thread pool
        loop = asyncio.get_event_loop()
        splitter = self._get_splitter(self.config.chunking_strategy)

        chunks = await loop.run_in_executor(
            None,
            splitter.split_text,
            text
        )

        # Create chunk objects
        result = []
        char_position = 0

        for idx, chunk_text in enumerate(chunks):
            chunk_id = self._generate_chunk_id(document_id, idx)

            # Find chunk position in original text
            char_start = text.find(chunk_text, char_position)
            if char_start == -1:
                char_start = char_position
            char_end = char_start + len(chunk_text)
            char_position = char_end

            result.append(DocumentChunk(
                chunk_id=chunk_id,
                document_id=document_id,
                content=chunk_text,
                chunk_index=idx,
                total_chunks=len(chunks),
                metadata={
                    **metadata,
                    "chunk_index": idx,
                    "total_chunks": len(chunks),
                },
                char_start=char_start,
                char_end=char_end,
            ))

        logger.debug(
            "Chunks generated",
            extra={
                "document_id": document_id,
                "text_length": len(text),
                "chunks": len(result),
            }
        )

        return result

    @staticmethod
    def _generate_chunk_id(document_id: str, chunk_index: int) -> str:
        """Generate unique chunk ID."""
        combined = f"{document_id}:chunk:{chunk_index}"
        return hashlib.sha256(combined.encode()).hexdigest()[:16]


class EnhancedDocumentProcessor:
    """Production-quality document processor with streaming and async processing.

    Features:
    - Async file I/O with aiofiles
    - Streaming document processing
    - Multiple file format support
    - Efficient concurrent processing
    - Proper resource cleanup
    - Progress tracking
    - Error handling and recovery
    - Metadata extraction
    """

    def __init__(self, config: Optional[ProcessingConfig] = None):
        """Initialize document processor.

        Args:
            config: Processing configuration.
        """
        self.config = config or ProcessingConfig()
        self._text_extractor = TextExtractor()
        self._chunk_generator = ChunkGenerator(self.config)
        self._semaphore = asyncio.Semaphore(self.config.max_concurrent_files)

        logger.info(
            "Initializing document processor",
            extra={
                "chunk_size": self.config.chunk_size,
                "chunk_overlap": self.config.chunk_overlap,
                "strategy": self.config.chunking_strategy.value,
                "max_concurrent": self.config.max_concurrent_files,
            }
        )

    async def process_file(
        self,
        file_path: Path,
    ) -> ProcessedDocument:
        """Process single file into chunks.

        Args:
            file_path: Path to file.

        Returns:
            Processed document with chunks.
        """
        start_time = time.time()
        file_path = Path(file_path)

        # Validate file
        if not file_path.exists():
            return ProcessedDocument(
                document_id="",
                file_path=str(file_path),
                chunks=[],
                metadata={},
                processing_time=0,
                error=f"File not found: {file_path}"
            )

        if file_path.suffix.lower() not in self.config.supported_extensions:
            return ProcessedDocument(
                document_id="",
                file_path=str(file_path),
                chunks=[],
                metadata={},
                processing_time=0,
                error=f"Unsupported file type: {file_path.suffix}"
            )

        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > self.config.max_file_size_mb:
            return ProcessedDocument(
                document_id="",
                file_path=str(file_path),
                chunks=[],
                metadata={},
                processing_time=0,
                error=f"File too large: {file_size_mb:.2f}MB (max {self.config.max_file_size_mb}MB)"
            )

        # Generate document ID
        document_id = self._generate_document_id(file_path)

        # Extract metadata
        metadata = await self._extract_metadata(file_path) if self.config.enable_metadata_extraction else {}

        try:
            async with self._semaphore:
                # Extract text
                text = await self._text_extractor.extract_text(file_path)

                if not text.strip():
                    return ProcessedDocument(
                        document_id=document_id,
                        file_path=str(file_path),
                        chunks=[],
                        metadata=metadata,
                        processing_time=time.time() - start_time,
                        error="Empty document after text extraction"
                    )

                # Generate chunks
                chunks = await self._chunk_generator.generate_chunks(
                    text=text,
                    document_id=document_id,
                    metadata=metadata,
                )

                processing_time = time.time() - start_time

                logger.info(
                    "File processed successfully",
                    extra={
                        "file": file_path.name,
                        "document_id": document_id,
                        "chunks": len(chunks),
                        "duration_s": round(processing_time, 2),
                    }
                )

                return ProcessedDocument(
                    document_id=document_id,
                    file_path=str(file_path),
                    chunks=chunks,
                    metadata=metadata,
                    processing_time=processing_time,
                )

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(
                "File processing failed",
                extra={
                    "file": file_path.name,
                    "error": str(e),
                    "duration_s": round(processing_time, 2),
                }
            )

            return ProcessedDocument(
                document_id=document_id,
                file_path=str(file_path),
                chunks=[],
                metadata=metadata,
                processing_time=processing_time,
                error=str(e)
            )

    async def process_directory(
        self,
        directory_path: Path,
        recursive: bool = True,
    ) -> AsyncIterator[ProcessedDocument]:
        """Process all files in directory with streaming output.

        Args:
            directory_path: Path to directory.
            recursive: Whether to process subdirectories.

        Yields:
            Processed documents as they complete.
        """
        directory_path = Path(directory_path)

        if not directory_path.is_dir():
            raise ValueError(f"Not a directory: {directory_path}")

        # Collect all files
        pattern = "**/*" if recursive else "*"
        files = [
            f for f in directory_path.glob(pattern)
            if f.is_file() and f.suffix.lower() in self.config.supported_extensions
        ]

        logger.info(
            "Processing directory",
            extra={
                "directory": str(directory_path),
                "files": len(files),
                "recursive": recursive,
            }
        )

        # Process files with streaming output
        async for doc in self.process_files_stream(files):
            yield doc

    async def process_files_stream(
        self,
        file_paths: List[Path],
    ) -> AsyncIterator[ProcessedDocument]:
        """Process multiple files with streaming output.

        Args:
            file_paths: List of file paths.

        Yields:
            Processed documents as they complete.
        """
        if not file_paths:
            return

        # Create queue for results
        queue: asyncio.Queue[Optional[ProcessedDocument]] = asyncio.Queue(
            maxsize=self.config.stream_buffer_size
        )

        async def producer():
            """Process files and put results in queue."""
            try:
                tasks = [self.process_file(path) for path in file_paths]

                # Process with limit on concurrency
                for coro in asyncio.as_completed(tasks):
                    try:
                        doc = await coro
                        await queue.put(doc)
                    except Exception as e:
                        logger.error(f"Processing task failed: {e}")

                # Signal completion
                await queue.put(None)

            except Exception as e:
                logger.error(f"Producer failed: {e}")
                await queue.put(None)

        # Start producer
        producer_task = asyncio.create_task(producer())

        # Yield results as they arrive
        processed_count = 0
        try:
            while True:
                doc = await queue.get()
                if doc is None:
                    break

                processed_count += 1
                yield doc

                if processed_count % 10 == 0:
                    logger.info(
                        "Processing progress",
                        extra={
                            "processed": processed_count,
                            "total": len(file_paths),
                        }
                    )

        finally:
            # Ensure producer completes
            if not producer_task.done():
                producer_task.cancel()
                try:
                    await producer_task
                except asyncio.CancelledError:
                    pass

        logger.info(
            "Stream processing completed",
            extra={
                "total_files": len(file_paths),
                "processed": processed_count,
            }
        )

    async def process_files_batch(
        self,
        file_paths: List[Path],
    ) -> List[ProcessedDocument]:
        """Process multiple files and return all results.

        Args:
            file_paths: List of file paths.

        Returns:
            List of processed documents.
        """
        if not file_paths:
            return []

        start_time = time.time()

        # Process all files concurrently
        tasks = [self.process_file(path) for path in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Handle exceptions
        processed_docs = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(
                    "File processing failed in batch",
                    extra={
                        "file": str(file_paths[i]),
                        "error": str(result),
                    }
                )
                # Create error document
                processed_docs.append(ProcessedDocument(
                    document_id="",
                    file_path=str(file_paths[i]),
                    chunks=[],
                    metadata={},
                    processing_time=0,
                    error=str(result)
                ))
            else:
                processed_docs.append(result)

        duration = time.time() - start_time
        successful = sum(1 for doc in processed_docs if doc.error is None)

        logger.info(
            "Batch processing completed",
            extra={
                "total_files": len(file_paths),
                "successful": successful,
                "failed": len(file_paths) - successful,
                "duration_s": round(duration, 2),
            }
        )

        return processed_docs

    @staticmethod
    def _generate_document_id(file_path: Path) -> str:
        """Generate unique document ID from file path."""
        # Use absolute path for consistency
        abs_path = str(file_path.absolute())
        return hashlib.sha256(abs_path.encode()).hexdigest()[:16]

    @staticmethod
    async def _extract_metadata(file_path: Path) -> Dict[str, Any]:
        """Extract file metadata.

        Args:
            file_path: Path to file.

        Returns:
            Metadata dictionary.
        """
        stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))

        return {
            "file_name": file_path.name,
            "file_path": str(file_path.absolute()),
            "file_extension": file_path.suffix.lower(),
            "file_size_bytes": stat.st_size,
            "mime_type": mime_type or "application/octet-stream",
            "created_at": stat.st_ctime,
            "modified_at": stat.st_mtime,
        }

    def get_stats(self) -> Dict[str, Any]:
        """Get processor statistics.

        Returns:
            Statistics dictionary.
        """
        return {
            "chunk_size": self.config.chunk_size,
            "chunk_overlap": self.config.chunk_overlap,
            "strategy": self.config.chunking_strategy.value,
            "max_concurrent": self.config.max_concurrent_files,
            "supported_extensions": list(self.config.supported_extensions),
        }
